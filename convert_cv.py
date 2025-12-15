#!/usr/bin/env python3
"""
Simple HTML to DOCX converter for Anonymous CV
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# Set A4 margins
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# Read HTML
with open('Anonymous_CV.html', 'r', encoding='utf-8') as f:
    content = f.read()

# Remove style tags and extract body content
body_match = re.search(r'<body>(.*?)</body>', content, re.DOTALL)
if body_match:
    body_content = body_match.group(1)
else:
    body_content = content

# Simple text extraction and formatting
lines = body_content.split('\n')
current_para = None
in_list = False

for line in lines:
    line = line.strip()
    if not line or line.startswith('<!--'):
        continue
    
    # Title
    if '<h1>' in line:
        match = re.search(r'<h1>(.*?)</h1>', line)
        if match:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(match.group(1).strip())
            run.font.size = Pt(12)
            run.font.bold = True
            para.paragraph_format.space_after = Pt(3)
    
    # Section headers
    elif '<h2>' in line:
        match = re.search(r'<h2>(.*?)</h2>', line)
        if match:
            para = doc.add_paragraph()
            run = para.add_run(match.group(1).strip())
            run.font.size = Pt(11)
            run.font.bold = True
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)
    
    # Job titles / h3
    elif '<h3>' in line:
        match = re.search(r'<h3>(.*?)</h3>', line)
        if match:
            para = doc.add_paragraph()
            run = para.add_run(match.group(1).strip())
            run.font.size = Pt(10)
            run.font.bold = True
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(2)
    
    # Paragraphs
    elif '<p' in line and '</p>' in line:
        match = re.search(r'<p[^>]*>(.*?)</p>', line, re.DOTALL)
        if match:
            text = re.sub(r'<[^>]+>', '', match.group(1))
            text = text.replace('&nbsp;', ' ').strip()
            if text:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(10)
                para.paragraph_format.space_after = Pt(2)
    
    # List items
    elif '<li>' in line:
        match = re.search(r'<li>(.*?)</li>', line, re.DOTALL)
        if match:
            text = re.sub(r'<[^>]+>', '', match.group(1))
            text = text.replace('&nbsp;', ' ').strip()
            if text:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(text)
                run.font.size = Pt(10)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.left_indent = Cm(0.5)

doc.save('Anonymous_CV.docx')
print("Successfully created Anonymous_CV.docx")

