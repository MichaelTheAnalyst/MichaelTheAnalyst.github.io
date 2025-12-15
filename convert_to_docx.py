#!/usr/bin/env python3
"""
Convert Anonymous_CV.html to Anonymous_CV.docx
"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Create a new Document
doc = Document()

# Set page margins (A4 size)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# Read the HTML file
with open('Anonymous_CV.html', 'r', encoding='utf-8') as f:
    html_content = f.read()

# Extract body content (remove HTML tags but keep structure)
from html.parser import HTMLParser

class CVParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.doc = Document()
        self.current_para = None
        self.in_list = False
        self.list_items = []
        self.in_li = False
        self.current_text = ""
        
    def handle_starttag(self, tag, attrs):
        if tag == 'h1':
            self.current_para = self.doc.add_paragraph()
            self.current_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = self.current_para.add_run()
            run.font.size = Pt(12)
            run.font.bold = True
        elif tag == 'h2':
            if self.current_para:
                self.current_para.add_run().add_break()
            self.current_para = self.doc.add_paragraph()
            run = self.current_para.add_run()
            run.font.size = Pt(11)
            run.font.bold = True
            # Add border effect
            self.current_para.paragraph_format.space_before = Pt(6)
            self.current_para.paragraph_format.space_after = Pt(3)
        elif tag == 'h3':
            if self.current_para:
                self.current_para.add_run().add_break()
            self.current_para = self.doc.add_paragraph()
            run = self.current_para.add_run()
            run.font.size = Pt(10)
            run.font.bold = True
            self.current_para.paragraph_format.space_before = Pt(4)
            self.current_para.paragraph_format.space_after = Pt(2)
        elif tag == 'p':
            if not self.in_list:
                self.current_para = self.doc.add_paragraph()
                self.current_para.paragraph_format.space_after = Pt(2)
        elif tag == 'ul':
            self.in_list = True
            self.list_items = []
        elif tag == 'li':
            self.in_li = True
            self.current_text = ""
            
    def handle_endtag(self, tag):
        if tag == 'h1' or tag == 'h2' or tag == 'h3':
            self.current_para = None
        elif tag == 'p':
            if self.current_para and self.current_text:
                self.current_para.add_run(self.current_text.strip())
                self.current_text = ""
        elif tag == 'ul':
            if self.list_items:
                for item in self.list_items:
                    para = self.doc.add_paragraph(style='List Bullet')
                    para.paragraph_format.space_after = Pt(1)
                    run = para.add_run(item.strip())
                    run.font.size = Pt(10)
            self.in_list = False
            self.list_items = []
        elif tag == 'li':
            if self.current_text:
                self.list_items.append(self.current_text)
            self.in_li = False
            self.current_text = ""
            
    def handle_data(self, data):
        if self.in_li:
            self.current_text += data
        elif self.current_para:
            if not self.current_para.runs:
                run = self.current_para.add_run()
                run.font.size = Pt(10)
            self.current_para.runs[-1].text += data

# Parse HTML
parser = CVParser()
parser.feed(html_content)

# Save as docx
parser.doc.save('Anonymous_CV.docx')
print("Successfully created Anonymous_CV.docx")


